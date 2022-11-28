# streamlit run web.py --global.dataFrameSerialization="legacy"
import pandas as pd
import streamlit as st
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode,  DataReturnMode, JsCode
import all_function as af
from docx import Document
import io

st.sidebar.image("logo.png", width=200)

#st.sidebar.title("Content")
status_select = st.sidebar.selectbox(
        "Select page",
        ("Home", "Record")
    )

#st.set_page_config(page_title='Biohub: Publication & Preprint',layout='wide')

if status_select == 'Home':

    st.header('Biohub publication search result.')

    # add_selectbox = st.sidebar.selectbox(
    #     "page selection",
    #     ("record", "report")
    # )

    ##### 0. load data
    base = pd.read_csv('database/basedb.csv', encoding='utf-8-sig')
    allchangedb_old= pd.read_csv( 'database/changedb (old version).csv', encoding='utf-8-sig')
    allchangedb_new= pd.read_csv('database/changedb (new version).csv' , encoding='utf-8-sig')
    alldeletedb= pd.read_csv('database/deletedb.csv' , encoding='utf-8-sig')
    pub_pre= pd.read_csv('database/matched pub-preprint.csv',encoding='utf-8-sig')

    preprint_list=['biorxiv','bioRxiv','medrxiv','medRxiv','arxiv','arXiv']
    nopre=base[ ~(base['journal'].isin(preprint_list)) & base[['confirm preprint doi','possible match result']].notnull().any(1)]

    ##### 1. table selection
    def table_select(table_option):
        if table_option == 'basedb':
            df = base
        if table_option == 'changedb (old version)':
            df = allchangedb_old
        if table_option == 'changedb (new version)':
            df = allchangedb_new
        if table_option == 'matched pub-preprint':
            df = pub_pre
        if table_option == 'deletedb':
            df = alldeletedb
        if table_option == 'publication non-preprint':
            df=nopre
        df.fillna('', inplace=True)
        #df=af.standardize_name(df)
        #df.reset_index(drop=True,inplace=True)
        return df


    left,right=st.columns(2)
    with left:
        table_option = st.selectbox(
            'Which table you would like to check?',
            #('matched pub-preprint', 'basedb', 'changedb (old version)', 'changedb (new version)', 'deletedb'))
            ('basedb', 'changedb (old version)', 'changedb (new version)', 'matched pub-preprint','deletedb','publication non-preprint'))
    with right:
        format_select = st.radio(
        "What's table format you want?",
        ('short format','full text'))

    df=table_select(table_option)

    col1, col2 = st.columns(2)
    with col1:
        #### add search bar
        search = st.text_input('Enter search words: (Case Ignore)')

        if search:
            ind = []
            for i in df.columns:
                ind += list(df.loc[df[i].astype(str).str.lower()
                            .str.contains(search.lower())].index)
            df = df.iloc[list(set(ind)), :]
            reload_data = True

    #####
    gb = GridOptionsBuilder.from_dataframe(df)
    gb.configure_pagination(enabled=True, paginationAutoPageSize=True, paginationPageSize=20)
    if format_select == 'short format':
        gb.configure_default_column(editable=True, groupable=True, wrapText=True, enableRowGroup=True,
                                    aggFunc='sum', sizeColumnsToFit=True) 

    if format_select == 'full text':
        gb.configure_default_column(editable=True, groupable=True, wrapText=True, enableRowGroup=True,
                                    aggFunc='sum', sizeColumnsToFit=True, autoHeight=True)
        gb.configure_column(field="authors", maxWidth=200)  # if want to set for single row

    #gb.configure_columns(column_names=df.columns, maxWidth=200)  # column_names=[]
    gb.configure_selection(selection_mode='multiple', use_checkbox=True)

    grid_options = gb.build()
    grid_response = AgGrid(
        df,
        grid_options,
        update_mode=GridUpdateMode.VALUE_CHANGED | GridUpdateMode.SELECTION_CHANGED,
        enable_enterprise_modules=True
    )

    with col2:
        #### edit button
        st.write('Press **return/enter** after you change data.\n\n  If you confirm to edit data, click button below.')
        edit = st.button('Confirm edit!')

        if edit:
            df = grid_response['data']
            
            if table_option=='publication non-preprint':
                df.to_csv('database/nopre.csv', encoding='utf-8-sig', index=False)
            else:
                df.to_csv('database/'+table_option+'.csv', encoding='utf-8-sig', index=False)
                
            st.write('Edit successful!')
            

    #### 2) compare & detail function
    st.subheader('Detail & Compare Records')

    sel_row = grid_response['selected_rows']  # type: list

    def compre_sel(sel_df):
        sel_df.columns = sel_df.loc['title']
        sel_df = sel_df.drop(['save datetime', 'title'])
        sel_df.reset_index(inplace=True)
        #sel_df.reset_index(drop=True, inplace=True)

        gb_sel = GridOptionsBuilder.from_dataframe(sel_df)
        gb_sel.configure_column('index',  pinned='left',weight=20)  ###
        gb_sel.configure_default_column(autoHeight=True, groupable=True,
                                        wrapText=True,  value=True, enableRowGroup=True, aggFunc='sum')
        gb_sel.configure_columns(
            column_names=sel_df.columns, maxWidth=1700/sel_df.shape[1])
        grid_options_sel = gb_sel.build()
        grid_table_sel = AgGrid(sel_df, grid_options_sel, update_mode=GridUpdateMode.VALUE_CHANGED | GridUpdateMode.SELECTION_CHANGED,
                                enable_enterprise_modules=True)  # fit_columns_on_grid_load=True


    if len(sel_row) == 1:
        st.json(sel_row[0])

        match_result= pd.DataFrame(sel_row)['match result'].values[0]
        if match_result !='':
            dmatch=pd.DataFrame()
            search_dois = match_result.split(' ; ')
            for j in search_dois:
                search_doi=j.split('10.')[-1]
                dmatch=pd.concat([dmatch,base.loc[base.loc[:,'doi'].str.contains(search_doi)]])

            st.subheader('Possible Match Result')
            dmatch=dmatch.transpose()
            compre_sel(dmatch)

            m = dmatch.shape[1]
            if m==2 or m==3:
                if m==2:
                    col1, col2, col3 = st.columns((1, 3, 3))
                if m==3:
                    col1, col2, col3, col4 = st.columns((1, 2, 2,2))
                    button_phold3 = col3.empty()
                    do_action3 = button_phold1.button('Change to this', key=3)

                button_phold1 = col2.empty()
                do_action1 = button_phold1.button('Change to this', key=1)

                button_phold2 = col3.empty()
                do_action2 = button_phold2.button('Change to this', key=2)

                def change(do_action, i):
                    if do_action:
                        x=pd.DataFrame(sel_row)['doi'].values[0]

                        base.loc[base.loc[:, 'doi'].str.contains(x), :].to_csv( 'database/changedb (old version).csv',  mode='a', index=False, header=False, encoding='utf-8-sig')

                        base.loc[base.loc[:, 'doi'].str.contains(x), 'match result']=dmatch.loc['journal',dmatch.columns.values[i-1]]+' '+dmatch.loc['doi',dmatch.columns.values[i-1]]
                        base.to_csv('database/basedb.csv', encoding='utf-8-sig', index=False)

                        base.loc[base.loc[:, 'doi'].str.contains(x), :].to_csv('database/changedb (new version).csv' , mode='a', index=False, header=False,  encoding='utf-8-sig')

                        pub_pre.loc[pub_pre.loc[:, 'doi'].str.contains(x), 'match result']=dmatch.loc['journal',dmatch.columns.values[i-1]]+' '+dmatch.loc['doi',dmatch.columns.values[i-1]]
                        pub_pre.to_csv('database/matched pub-preprint.csv',encoding='utf-8-sig', index=False)
                        
                        nopre=base[ ~(base['journal'].isin(preprint_list)) & (base['match result'].notnull())] 
                        nopre.to_csv('database/nopre.csv', encoding='utf-8-sig', index=False)

                        st.write('Done')

                        import time
                        time.sleep(1)

                        pyautogui.hotkey("ctrl","F5")

                change(do_action1, 1)
                change(do_action2, 2)

                if m==3:
                    change(do_action3, 3)

            if m>=4:
                st.write('sorry please change manually.')


            dmatch = pd.DataFrame()


    if len(sel_row) > 1:
        sel_df = pd.DataFrame(sel_row).transpose()
        compre_sel(sel_df)


if status_select =='Record':

    st.header('Publication Report')

    col1, col2 = st.columns(2)
    with col1:
        start_date = st.date_input("Input Start Date", value=pd.to_datetime("2022-06-01"))
    with col2:
        end_date = st.date_input("Input End Date", value=pd.to_datetime("today"))
    
    
    start = start_date.strftime("%Y-%m-%d")
    end = end_date.strftime("%Y-%m-%d")

    st.write('The time period you choose is: '+str(start_date)+' ~ '+str(end_date))

    ##### coding
    # 1. Select author list 
    author=pd.read_csv('database/Biohub authors.csv', encoding='utf-8-sig')  #  biohub author
    author['Team']=author['Team'].fillna('NA')
    author=author.fillna('')
    fauthor_list=author['MatchName'].str.replace(',','').to_list()

    with st.expander("advance filter"):
        cola, colb = st.columns(2)
        with cola:
            sel_campus_sim= st.multiselect('Campus (simple)',
                ['All','UCSF', 'Stanford', 'Berkeley', 'Biohub'])
        with colb:
            sel_campus= st.multiselect('Campus',
                ['All','UCSF','UCSF/Gladstone', 'Stanford', 'Berkeley','Berkeley/LBNL', 'Biohub'])
            
        colc, cold = st.columns(2)
        with colc:
            sel_team= st.multiselect('team',
                ['All','NA', 'Computational Microscopy', 'Infectious Disease', 'Genome Engineering', 'Data Science', 'Proteomics/Mass Spec', 'Genomics', 'BioEngineering', 'Quantitative Cell Science'])
        with cold:
            sel_status= st.multiselect('status',
                ['All','Alumni', 'Former', 'Current'])
        
        sel_role= st.multiselect('role',
                ['All','Collaborator', 'Group leader', 'Investigator', 'Lab manager', 'Platform leader', 'President', 'Project Leader', 'Senior scientist/engineer'])
      
        click=st.button('submit')
        
        if click:  # using function to ahcieve this
            if sel_campus_sim==[] or sel_campus_sim==['All']:  
                con_campus_sim=['UCSF', 'Stanford', 'Berkeley', 'Biohub']
            else:
                con_campus_sim=sel_campus_sim
                
            if sel_campus==[] or sel_campus==['All']:
                con_campus=['UCSF','UCSF/Gladstone', 'Stanford', 'Berkeley','Berkeley/LBNL', 'Biohub']
            else:
                con_campus=sel_campus
                
            if sel_team==[] or sel_team==['All']:
                con_team=['NA', 'Computational Microscopy', 'Infectious Disease', 'Genome Engineering', 'Data Science', 'Proteomics/Mass Spec', 'Genomics', 'BioEngineering', 'Quantitative Cell Science']
            else:
                con_team=sel_team
                
            if sel_status==[] or sel_status==['All']:
                con_status=['Alumni', 'Former', 'Current']
            else:
                con_status=sel_status
                
            if sel_role==[] or sel_role==['All']:
                con_role=['Collaborator', 'Group leader', 'Investigator', 'Lab manager', 'Platform leader', 'President', 'Project Leader', 'Senior scientist/engineer']
            else:
                con_role=sel_role

            filter= (author['Campus (simple)'].isin(con_campus_sim) )&  (author['Campus'].isin(con_campus) )& (author['Role'].isin(con_role) )  & (author['Status'].isin(con_status) ) & (author['Team'].isin(con_team) )
            
            fauthor=author[filter]  # author table after filter 
            fauthor_list=fauthor['MatchName'].str.replace(',','').to_list()          

    # intramural=author[author['Campus (simple)'] == 'Biohub']
    # investigator=author.loc[author['Role'] == 'Investigator']
    # df_a=author.loc[(author['Campus (simple)'] == 'Biohub') | (author['Role'] == 'Investigator') ]
    biohub_staff_author=author[author['Campus (simple)'] == 'Biohub']['MatchName'].str.replace(',','').to_list()

    # 2. Choose publication & prerpint within specific time period 
    df = pd.read_csv('database/basedb.csv', encoding='utf-8-sig')
    
    # need to check date. many records have been filtered. don't know why yet :(
    df['date'] = pd.to_datetime(df['date'])  
    condition = (df['date'] >= start) & (df['date'] <= end) 
    df=df.loc[condition]
    
    st.write(df.shape)
    
    def format_col(df,col_name):
        # format column and replace original
        df=df.fillna('').reset_index(drop=True)
        
        # format 'corresponding author' column and replace original
        for ind,i in enumerate(df[col_name]):
            if i == '':
                continue 
            i_list=i.split(';')
            res=[]
            for j in i_list:  
                old_name,new_name,prob=af.authormatch(j)
                res.append(new_name)
            df.loc[ind,col_name] = '; '.join(m for m in res)
        df.fillna('', inplace=True)
        return df
    
    df=format_col(df,col_name='corresponding author') 
    df=format_col(df,col_name='biohub author')
    
    #df['b_au']=df[['biohub author','possible biohub author','format biohub author']].agg('; '.join, axis=1)#.str.replace('; ; ; ','; ').str.replace('; ; ','; ').str.strip('; ')
    #df['b_au']=df[['possible biohub author','possible biohub author','format biohub author','corresponding author']].agg('; '.join, axis=1)
    df['b_au']=df['biohub author']  # unique of b_au

    ind_list=[]
    biohub_staff_ind=[]
    for ind,i in enumerate(df['b_au']):
        i='; '.join(set([j.strip() for j in i.split(';') if j]))
        df.loc[ind,'b_au']=i
        for j in i.split(';'): 
            if j.strip(' ').replace(', ',' ') in fauthor_list:  # only return people who is in filter list
                # Problem: if there are some author not in 'biohub author.xlsx', it will not show theirs publications
                ind_list.append(ind)
            if j.strip(' ').replace(', ',' ') in biohub_staff_author:
                biohub_staff_ind.append(ind)
                
    ## remove duplicate index in list
    ind_list=[*set(ind_list)]
    biohub_staff_ind=[*set(biohub_staff_ind)]
        
    df=df.iloc[ind_list,:]

    # 2.2 divided into two categories: publication and preprint #得把date格式改一下然后换成date
    
    preprint_list=['biorxiv','bioRxiv','medrxiv','medRxiv','arxiv','arXiv']
    preprint = df[df['journal'].isin(preprint_list)]
    publication = df[~df['journal'].isin(preprint_list)]

    pre_condition = (df['journal'].isin(preprint_list))
    p1_pre=df.loc[pre_condition]

    pub_condition = (~df['journal'].isin(preprint_list))
    p1_pub=df.loc[pub_condition]
    
    ########
    tab1, tab2, tab3 = st.tabs(["summary report", "summary table", "author table"])
    
    with tab1:
        p1 ='##### Papers (Research articles, methods papers, reviews, etc.) and conference proceedings:\n  '

        doc = Document()
        title = doc.add_heading("Biohub intramural research program – Papers published and preprints first-deposited",0)
        #para = doc.add_paragraph("Includes papers, conference proceedings, and preprints published or first-deposited since the last Biohub All-Hands meeting that cite Biohub affiliation or funding and that include a Biohub employee as a co-author (we may easily have missed something, so please feel free to send Bill Burkholder any additions or corrections)\n")
        para = doc.add_paragraph()
        para.add_run("Includes papers, conference proceedings, and preprints published or first-deposited since the last Biohub All-Hands meeting that cite Biohub affiliation or funding and that include a Biohub employee as a co-author (we may easily have missed something, so please feel free to send Bill Burkholder any additions or corrections)\n").italic = True
        
        title = doc.add_heading('Papers (Research articles, methods papers, reviews, etc.) and conference proceedings\n\n  ', 2)  
        for ind,row in p1_pub.iterrows():
            intro=row['b_au']+" of "+row['corresponding author']+"’s lab at "+', '.join(row['corresponding author institution'].split(',')[0:2])
            detail=row['title']+' PMID: '+str(row['pmid'])+'\n'
            if row['corresponding author'] !='':
                p1 +='- **'+intro+'**\n\n    '+detail
                para = doc.add_paragraph()
                para.add_run(intro).bold = True
                para.add_run('\n\n      '+detail)  
            else:
                p1 +='- **'+row['b_au']+'**\n\n    '+detail
                para = doc.add_paragraph()
                para.add_run(row['b_au']).bold = True
                para.add_run('\n\n      '+detail)

        title = doc.add_heading("Preprints\n\n  ",2)
        p1 += "\n\n##### Preprints\n"
        for ind,row in p1_pre.iterrows():
            intro=row['b_au']+" of "+row['corresponding author']+"’s lab at "+', '.join(row['corresponding author institution'].split(',')[0:2])
            detail2=row['title']+' doi: '+str(row['doi'])+'\n'
            if row['corresponding author'] !='':
                p1 +='- **'+intro+'**\n\n    '+detail2
                para = doc.add_paragraph()
                para.add_run(intro).bold = True
                para.add_run('\n\n      '+detail2)  
            else:
                p1 +='- **'+row['b_au']+'**\n\n    '+detail2
                para = doc.add_paragraph()
                para.add_run(row['b_au']).bold = True
                para.add_run('\n\n      '+detail2)
    
        p1=p1.replace('; ; ','; ')
        st.markdown("### Biohub intramural research program – Papers published and preprints first-deposited \n\nIncludes papers, conference proceedings, and preprints published or first-deposited since the last Biohub All-Hands meeting that cite Biohub affiliation or funding and that include a Biohub employee as a co-author (we may easily have missed something, so please feel free to send Bill Burkholder any additions or corrections)\n")
        
        col1, col2 = st.columns([3,2],gap = "small")
        with col1:
            bio = io.BytesIO()
            doc.save(bio)
            if doc:
                st.download_button(label="Download docx version",data=bio.getvalue(),
                    file_name="Report.docx",mime="docx")
        
        with col2:    
            st.download_button(label="Download md version",data=p1,file_name="Report.md",)
        st.markdown(p1)  

    with tab2:    
        df2=pd.DataFrame(columns =['Biohub staff authors', 'All authors'],index = ['Papers (Research articles, methods papers, reviews, etc.) and conference proceedings', 'Preprints', 'Total'])

        df2.iloc[0,1]= len(p1_pub.index)
        df2.iloc[1,1]= len(p1_pre.index)
        df2.iloc[2,1]=df2.iloc[0,1]+df2.iloc[1,1]

        df2.iloc[0,0]=len([i for i in biohub_staff_ind if i in p1_pub.index])
        df2.iloc[1,0]=len([i for i in biohub_staff_ind if i in p1_pre.index])

        df2.iloc[2,0]=df2.iloc[0,0]+df2.iloc[1,0]

        st.markdown("### Papers published and preprints first-deposited \n")
        st.download_button(label='Download Report',data=df2.to_csv().encode('utf-8'),file_name='Overall Report.csv')
        st.write(df2)
             
    with tab3:
        p3=pd.DataFrame(index = author['MatchName'])
        p3.index.name = None

        # column 1
        c1=df.loc[df['corresponding author']!='',['title','journal','corresponding author']].set_index(['title','journal'])['corresponding author'].str.split("; ", expand=True).stack().reset_index(drop=True, level=-1).reset_index().groupby([0]).size().rename("Total articles as corresponding author")
        p3=pd.merge(p3,c1,how='left',left_index=True,right_index=True)

        # column 2
            ## Exclude 'review list'
        with open('database/list of review journals.txt') as file:
            review_list = [line.rstrip() for line in file]

        df=df[~df['journal'].isin(review_list)]
        
        c2=df.loc[df['corresponding author']!='',['title','journal','corresponding author']].set_index(['title','journal'])['corresponding author'].str.split("; ", expand=True).stack().reset_index(drop=True, level=-1).reset_index().groupby([0]).size().rename("Qualifying articles as corresponding author")
        p3=pd.merge(p3,c2,how='left',left_index=True,right_index=True)

        # column 3
        c3=df.loc[  (df['corresponding author']!='') & (df['confirm preprint doi']!='') ,['title','journal','corresponding author']].set_index(['title','journal'])['corresponding author'].str.split("; ", expand=True).stack().reset_index(drop=True, level=-1).reset_index().groupby([0]).size().rename("Qualifying articles as corresponding author with preprints")
        p3=pd.merge(p3,c3,how='left',left_index=True,right_index=True)

        ## fill na value with 0
        p3.iloc[:,0:3]=p3.fillna(0).iloc[:,0:3].astype('Int64')
                
        # column 4
        p3['Compliance']=100*p3['Qualifying articles as corresponding author with preprints']/p3['Qualifying articles as corresponding author']
        
        st.markdown("### Individual Author Report \n")
        st.download_button(label='Download Report',data=p3.to_csv().encode('utf-8'),file_name='Individual Author Report.csv')
        st.write(p3)
        
